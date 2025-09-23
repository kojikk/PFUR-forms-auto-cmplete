from flask import Flask, render_template, request, send_file, flash, redirect, url_for
import os
import shutil
from datetime import datetime
from docx import Document
import tempfile
import io

app = Flask(__name__)
app.secret_key = 'your-secret-key-here'

# Путь к шаблону документа
TEMPLATE_PATH = 'template.docx'

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate_document():
    try:
        # Получаем данные из формы
        form_data = {
            'project_name': request.form.get('project_name', '').strip(),
            'key_words': request.form.get('key_words', '').strip(),
            'sience_field': request.form.get('sience_field', '').strip(),
            'research_direction': request.form.get('research_direction', '').strip(),
            'project_category': request.form.get('project_category', '').strip(),
            'project_annotation': request.form.get('project_annotation', '').strip(),
            'name_of_np': request.form.get('name_of_np', '').strip(),
            'head_of_project': request.form.get('head_of_project', '').strip(),
            'head_of_np': request.form.get('head_of_np', '').strip(),
            'date': request.form.get('date', '').strip()
        }
        
        # Проверяем, что все обязательные поля заполнены
        required_fields = ['project_name', 'key_words', 'sience_field', 'research_direction', 
                          'project_category', 'project_annotation', 'name_of_np', 
                          'head_of_project', 'head_of_np', 'date']
        
        missing_fields = [field for field in required_fields if not form_data[field]]
        if missing_fields:
            flash(f'Пожалуйста, заполните обязательные поля: {", ".join(missing_fields)}')
            return redirect(url_for('index'))
        
        # Преобразуем дату из формата YYYY-MM-DD в русский формат
        if form_data['date']:
            try:
                date_obj = datetime.strptime(form_data['date'], '%Y-%m-%d')
                months = ['января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
                         'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря']
                form_data['date'] = f"{date_obj.day} {months[date_obj.month-1]} {date_obj.year}г."
            except ValueError:
                flash('Неверный формат даты!')
                return redirect(url_for('index'))
        
        # Создаем копию шаблона
        if not os.path.exists(TEMPLATE_PATH):
            flash('Файл шаблона не найден!')
            return redirect(url_for('index'))
        
        # Создаем временный файл для обработки
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            shutil.copy2(TEMPLATE_PATH, temp_file.name)
            
            # Открываем документ и заменяем плейсхолдеры
            doc = Document(temp_file.name)
            
            # Заменяем плейсхолдеры в параграфах
            for paragraph in doc.paragraphs:
                for key, value in form_data.items():
                    placeholder = f'{{{key}}}'
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, value)
            
            # Заменяем плейсхолдеры в таблицах
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in form_data.items():
                            placeholder = f'{{{key}}}'
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, value)
            
            # Сохраняем изменения
            doc.save(temp_file.name)
            
            # Создаем имя файла для скачивания
            output_filename = f'Заявка {form_data["project_name"]}.docx'
            
            # Отправляем файл пользователю
            return send_file(
                temp_file.name,
                as_attachment=True,
                download_name=output_filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
    except Exception as e:
        flash(f'Произошла ошибка при генерации документа: {str(e)}')
        return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True)
